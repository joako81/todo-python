# Importa módulos necesarios para trabajar con sistemas operativos, fechas, tablas (Pandas), documentos (ReportLab) y colores.
#pip install pandas
#pip install python-docx
#pip install reportlab
import os
import datetime
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.colors import blue
from reportlab.lib.units import inch

# Definir una clase para representar una tarea individual.
class Tarea:
# Inicializa una tarea con una descripción, establece su estado inicial como "No comenzada",
# registra la fecha y hora actuales como inicio y fecha de cambio, y deja la fecha de fin y tiempo transcurrido vacíos.
    def __init__(self, descripcion):
        self.descripcion = descripcion.strip()
        self.estado = "No comenzada"
        self.fecha_inicio = datetime.datetime.now()
        self.fecha_fin = None
        self.fecha_cambio = datetime.datetime.now()
        self.tiempo_transcurrido = ""

# Actualiza el estado de la tarea y ajusta las fechas de inicio, fin y cambio según el nuevo estado.
    def actualizar_estado(self, nuevo_estado):
        self.estado = nuevo_estado
        if self.estado == "Realizando":
            self.fecha_inicio = datetime.datetime.now()
            self.fecha_cambio = self.fecha_inicio
        elif self.estado == "Completado":
            self.fecha_fin = datetime.datetime.now()
            self.fecha_cambio = self.fecha_fin
            self.calcular_tiempo_transcurrido()
            
# Calcula el tiempo transcurrido desde el inicio hasta el fin de la tarea, si ambos están definidos.
    def calcular_tiempo_transcurrido(self):
        if self.fecha_inicio and self.fecha_fin:
            diferencia_total = self.fecha_fin - self.fecha_inicio
            horas = diferencia_total.total_seconds() / 3600
            minutos = (diferencia_total.total_seconds() % 3600) // 60
            segundos = diferencia_total.total_seconds() % 60
            self.tiempo_transcurrido = f"{int(horas)} horas, {int(minutos)} minutos y {segundos} segundos"

# Retorna el estado de la tarea con un color correspondiente según el estado actual usando ANSI escape codes.
    def obtener_estado_con_color(self):
        if self.estado == "No comenzada":
            return "\033[94m{}\033[0m".format(self.estado)
        elif self.estado == "Realizando":
            return "\033[93m{}\033[0m".format(self.estado)
        elif self.estado == "Completado":
            return "\033[92m{}\033[0m".format(self.estado)
        else:
            return self.estado
 
 
 # Retorna la fecha y hora de último cambio de la tarea, o un mensaje indicando que no se registró.
    def obtener_fecha_cambio(self):
        if self.fecha_cambio:
            return self.fecha_cambio.strftime("%Y-%m-%d %H:%M:%S")
        else:
            return "Fecha del cambio no registrada"
        
    def editar_tareas(self, posicion, nueva_descripcion):
        if 0 <= posicion < len(self.tareas):
            tarea = self.tareas[posicion]
            tarea.descripcion = nueva_descripcion.strip()
            self.tareas.sort(key=lambda tarea: ('No comenzada', 'Realizando', 'Completado').index(tarea.estado))
            print("La tarea ha sido editada exitosamente.")
        else:
            print("\033[91mError: La posición ingresada es inválida.\033[0m")


# Definimos una clase para gestionar la lista de tareas.
class ListaDeTareas:
# Inicializa una lista vacía de tareas.
    def __init__(self):
        self.tareas = []
        
# Agrega tareas a la lista, validando que no sean cadenas vacías.
    def agregar_tareas(self, descripciones):
        for descripcion in descripciones:
            if not descripcion.strip():  # Verifica si la descripción está vacía punto de control para no introducir tareas en blanco
                print("\033[91mError: No se pueden agregar tareas vacías. Intente de nuevo.\033[0m")
                return  # Sale de la función si encuentra una tarea vacía
            tarea = Tarea(descripcion)
            self.tareas.append(tarea)

# Ordena y muestra las tareas por estado, coloreando el estado y mostrando detalles específicos dependiendo del estado.
    def mostrar_tareas(self):
        tareas_ordenadas = sorted(self.tareas, key=lambda tarea: ('No comenzada', 'Realizando', 'Completado').index(tarea.estado))
        if not tareas_ordenadas:
            print(" No hay tareas pendientes.")
        else:
            for i, tarea in enumerate(tareas_ordenadas, start=1):
                estado_coloreado = tarea.obtener_estado_con_color()
                if tarea.estado == "No comenzada":
                    tiempo_transcurrido = ""
                    fecha_creacion = tarea.fecha_cambio if tarea.fecha_cambio else "Fecha de creación no registrada"
                elif tarea.estado == "Realizando":
                    tiempo_transcurrido = ""
                    fecha_creacion = tarea.fecha_cambio
                elif tarea.estado == "Completado":
                    tiempo_transcurrido = tarea.tiempo_transcurrido
                    fecha_creacion = tarea.fecha_inicio
                    fecha_finalizacion = tarea.fecha_fin
                    info_completa = f". La tarea se completó el día: {fecha_finalizacion}, siendo el tiempo transcurrido en finalizar la tarea de: {tiempo_transcurrido}"
                print(f"{i}. {tarea.descripcion} ({estado_coloreado}) La fecha de creación de la tarea fue el día: {fecha_creacion}{info_completa if tarea.estado == 'Completado' else ''}")

# Funciones para guardar la lista de tareas en diferentes formatos de archivo.
# Guarda la lista de tareas en un archivo de texto.
    def guardar_tareas_en_txt(self, nombre_archivo):
        
        with open(nombre_archivo, 'w') as archivo:
            archivo.write("LISTA DE TUS TAREAS\n\n")
            for i, tarea in enumerate(self.tareas, start=1):
                archivo.write(f"{i}. {tarea.descripcion} ({tarea.estado})\n")

# Convierte la lista de tareas en un DataFrame de Pandas y guarda en un archivo Excel.

    def guardar_tareas_en_xls(self, nombre_archivo):
        datos_tareas = [
            {
                'Descripción': tarea.descripcion,
                'Estado': tarea.estado,
                'Fecha de Creación': tarea.fecha_cambio.strftime("%Y-%m-%d %H:%M:%S"),
                'Tiempo Transcurrido': tarea.tiempo_transcurrido if tarea.estado == "Completado" else "N/A"
            } for tarea in self.tareas
        ]
        
        df = pd.DataFrame(datos_tareas)
        df.to_excel(nombre_archivo , index=False)


# Crea un documento Word y agrega las tareas.
    def guardar_tareas_en_docx(self, nombre_archivo):
        doc = Document()
        doc.add_heading('Lista de Tus Tareas', level=1)
        for i, tarea in enumerate(self.tareas, start=1):
            doc.add_paragraph(f"{i}. {tarea.descripcion} ({tarea.estado})")
        doc.save(nombre_archivo)
        
# Genera un documento PDF con las tareas.
    def guardar_tareas_en_pdf(self, nombre_archivo):
        doc = SimpleDocTemplate(nombre_archivo, pagesize=letter)
        styles = getSampleStyleSheet()
        
        title_style = styles["Title"]
        title_style.textColor = blue
        title_style.fontName = "Helvetica-Bold"
        title_style.fontSize = 20
        
        story = [Paragraph("LISTA DE TUS TAREAS", title_style)]
        
        for i, tarea in enumerate(self.tareas, start=1):
            story.append(Paragraph(f"{i}. {tarea.descripcion} ({tarea.estado})"))
            story.append(Spacer(1, 12))
        
        doc.build(story)

    def editar_tareas(self, posicion, nueva_descripcion):
        if 0 <= posicion < len(self.tareas):
            tarea = self.tareas[posicion]
            tarea.descripcion = nueva_descripcion.strip()
            self.tareas.sort(key=lambda tarea: ('No comenzada', 'Realizando', 'Completado').index(tarea.estado))
            print("La tarea ha sido editada exitosamente.")
        else:
            print("\033[91mError: La posición ingresada es inválida.\033[0m")

# Función principal que ejecuta la aplicación.

def main():
    lista_de_tareas = ListaDeTareas()     # Inicializa la lista de tareas y presenta un menú interactivo para gestionarla.


    while True:
        # Menú de opciones para la gestión de tareas.
        # Las opciones incluyen agregar tareas, iniciar tareas, marcar tareas como completadas, editar tareas, guardar tareas en varios formatos de archivo y salir de la aplicación.
        # Utiliza funciones auxiliares para realizar acciones basadas en la elección del usuario.
        print("\033[34;1m***************************\n***************************\n**    APLICACIÓN PARA    **\n** GESTIONAR TUS TAREAS  **\n**                       **\n***************************\n***************************\033[0m")

        print("\n\033[95mLISTADO DE TAREAS\033[0m")

        lista_de_tareas.mostrar_tareas()
       
        print("\nMenú:")
        print("\033[93m1. Agregar tarea\033[0m")  
        print("\033[92m2. Iniciar tarea\033[0m")  
        print("\033[91m3. Marcar tarea como completada\033[0m")  
        print("\033[95m4. Editar tarea\033[0m")  
        print("\033[96m5. Guardar tareas en archivo\033[0m")  
        print("\033[97m6. Salir\033[0m")
        
        try:
            opcion = int(input("Seleccione una opción: "))
            if opcion == 1:
                while True:
                    try:
                        print("\033[93mIngrese las descripciones de las tareas separadas por comas: \033[0m")
                        descripciones = input().split(',')
                        lista_de_tareas.agregar_tareas(descripciones)
                        break  # Si no ocurre ninguna excepción, sale del bucle
                    except Exception as e:
                        print(f"\033[91mError: {str(e)}. Por favor, intente nuevamente.\033[0m")
                        seguir = input("\n¿Quieres agregar una tarea? o prefieres salir (s/n): ").lower()
                        if seguir!= 's':
                            break  # Salir del bucle si el usuario decide no volver a intentarlo
                lista_de_tareas.mostrar_tareas()  # Llamar a mostrar_tareas después de agregar tareas
                
            elif opcion == 2:
                lista_de_tareas.mostrar_tareas()
                while True:
                    print("\033[92mIngrese la posición de la tarea a iniciar (separado por comas): \033[0m")
                    posiciones_str = input()
                    posiciones = [int(pos) - 1 for pos in posiciones_str.split(',') if pos.isdigit()]

                    # Verificar si todas las posiciones son válidas
                    if all(0 <= pos < len(lista_de_tareas.tareas) for pos in posiciones):
                        for pos in posiciones:
                            tarea = lista_de_tareas.tareas[pos]
                            # Verificar si la tarea no está completada antes de permitir el cambio de estado
                            if tarea.estado!= "Completado":
                                tarea.actualizar_estado("Realizando")
                                print("La tarea ha sido iniciada exitosamente.")
                            else:
                                print("\033[91mLa tarea ya está completada. ¿Desea seleccionar otra tarea? (s/n)\033[0m")
                                seguir = input().lower()
                                if seguir == 's':
                                    continue  # Continuar con la siguiente iteración del bucle
                                else:
                                    break  # Salir del bucle si el usuario decide no continuar
                        break  # Salir del bucle si todas las posiciones son válidas y procesadas
                    else:
                        print("\033[91mError: Una o más posiciones ingresadas son inválidas. Intente de nuevo.\033[0m")

                    seguir = input("\n¿Desea volver a marcar un número para iniciar tarea? (s/n): ").lower()
                    if seguir!= 's':
                        break  # Salir del bucle externo si el usuario decide no volver a intentarlo
                lista_de_tareas.mostrar_tareas()  # Llamar a mostrar_tareas después de iniciar tareas
                
            elif opcion == 3:
                lista_de_tareas.mostrar_tareas()
                while True:
                    print("\033[91mIngrese la posición de la tarea a marcar como completada (separado por comas): \033[0m")
                    posiciones_str = input()
                    posiciones = [int(pos) - 1 for pos in posiciones_str.split(',') if pos.isdigit()]

                    # Verificar si todas las posiciones son válidas
                    if all(0 <= pos < len(lista_de_tareas.tareas) for pos in posiciones):
                        for pos in posiciones:
                            tarea = lista_de_tareas.tareas[pos]
                            tarea.actualizar_estado("Completado")
                            tarea.calcular_tiempo_transcurrido()
                        lista_de_tareas.mostrar_tareas()  # Muestra las tareas actualizadas
                        break  # Salir del bucle si todas las posiciones son válidas y procesadas
                    else:
                        print("\033[91mError: Una o más posiciones ingresadas son inválidas. Intente de nuevo.\033[0m")

                    seguir = input("\n¿Desea volver a intentar marcar un número para completar tarea? (s/n): ").lower()
                    if seguir!= 's':
                        break  # Salir del bucle externo si el usuario decide no volver a intentarlo

                
            elif opcion == 4:  # Opción para editar tareas
                lista_de_tareas.mostrar_tareas()
                while True:
                    print("\033[95mIngrese la posición de la tarea a editar (separado por comas): \033[0m")
                    posiciones_str = input()
                    posiciones = [int(pos) - 1 for pos in posiciones_str.split(',') if pos.isdigit()]

                    # Verificar si todas las posiciones son válidas
                    if all(0 <= pos < len(lista_de_tareas.tareas) for pos in posiciones):
                        for pos in posiciones:
                            tarea = lista_de_tareas.tareas[pos]
                            # Verificar si la tarea no está completada antes de permitir la edición
                            if tarea.estado!= "Completado":
                                print("Ingrese la nueva descripción de la tarea:")
                                nueva_descripcion = input()
                                tarea.descripcion = nueva_descripcion.strip()
                                lista_de_tareas.editar_tareas(pos, nueva_descripcion)  # Asumiendo que esta función está correctamente implementada
                                lista_de_tareas.mostrar_tareas()  # Muestra las tareas actualizadas
                            else:
                                print("\033[91mNo se puede editar una tarea que ya está completada.\033[0m")
                        break  # Salir del bucle si todas las posiciones son válidas y procesadas
                    else:
                        print("\033[91mError: Una o más posiciones ingresadas son inválidas. Intente de nuevo.\033[0m")

                    seguir = input("\¿Quieres elegir un número de tarea válido? (s/n): ").lower()
                    if seguir!= 's':
                        break  # Salir del bucle externo si el usuario decide no continuar editando
                lista_de_tareas.mostrar_tareas()  # Llamar a mostrar_tareas después de editar tareas

                
            elif opcion == 5:  # Guardar tareas en archivo
                solicitar_nombre_archivo = True
                while True:
                    print("\nSeleccione el formato de archivo para guardar las tareas:")
                    print("\033[95m1. TXT\033[0m")
                    print("\033[92m2. XLS\033[0m")
                    print("\033[96m3. DOCX\033[0m")
                    print("\033[93m4. PDF\033[0m")
                    print("\033[98m5. Salir sin guardar\033[0m")
                    try:
                        opcion = int(input("Seleccione una opción: "))
                        
                        if opcion == 5:
                            print("\nSaliendo de la aplicación sin guardar...\n")
                            solicitar_nombre_archivo = False  # Cambiar a False para no solicitar el nombre del archivo en futuras iteraciones
                            break  # Continuar con el siguiente ciclo del bucle, mostrando el menú nuevamente
                        
                        if solicitar_nombre_archivo:
                            nombre_archivo = input("Ingrese el nombre del archivo: ").strip()  # Limpiar el nombre del archivo
                            # Validar el nombre del archivo para eliminar espacios y caracteres no deseados
                            nombre_archivo_validado = ''.join(c for c in nombre_archivo if c.isalnum() or c in ['.', '-', '_'])
                        
                        if opcion == 1:
                            lista_de_tareas.guardar_tareas_en_txt(nombre_archivo_validado + ".txt")
                        elif opcion == 2:
                            lista_de_tareas.guardar_tareas_en_xls(nombre_archivo_validado + ".xlsx")
                        elif opcion == 3:
                            lista_de_tareas.guardar_tareas_en_docx(nombre_archivo_validado + ".docx")
                        elif opcion == 4:
                            lista_de_tareas.guardar_tareas_en_pdf(nombre_archivo_validado + ".pdf")
                        
                    except ValueError:
                        print("\033[91mError: Seleccione una opción válida.\033[0m")
                        
            elif opcion == 6:  # Salir
                print("Saliendo de la aplicación...")
                break
        except ValueError:
            print("\033[91mError: Seleccione una opción válida.\033[0m")

 # Ejecuta la función principal si este script se ejecuta directamente.
if __name__ == "__main__":
    main()